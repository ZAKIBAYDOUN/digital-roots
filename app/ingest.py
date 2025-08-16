import os, sys, json, hashlib, time, traceback, re, subprocess, shutil
from pathlib import Path
from typing import List, Dict, Any

import chromadb
from chromadb.config import Settings

# --- Optional: set TESSERACT_CMD if not on PATH
TESSERACT_CMD = os.getenv("TESSERACT_CMD", "tesseract")

# Deps per type
import fitz  # PyMuPDF
import docx
import pandas as pd
from PIL import Image
import pytesseract

# .doc converter (best effort)
try:
    import mammoth  # reads .doc via conversion to HTML/text
except Exception:
    mammoth = None

# ----------------- utils -----------------
sha = lambda b: hashlib.sha256(b).hexdigest()

def file_sig(p: Path) -> str:
    st = p.stat(); return sha(f"{p.resolve()}::{st.st_size}::{int(st.st_mtime)}".encode())

def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()

# ----------------- extractors -----------------

def extract_pdf(p: Path) -> List[str]:
    doc = fitz.open(p)
    out = []
    for page in doc:
        txt = page.get_text("text")
        if not norm(txt):  # if page empty, try OCR rasterization
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            try:
                img = Image.open(io.BytesIO(img_bytes))
            except Exception:
                from PIL import Image as PILImage
                import io as _io
                img = PILImage.open(_io.BytesIO(img_bytes))
            ocr = pytesseract.image_to_string(img)
            txt = ocr
        out.append(norm(txt))
    return out


def extract_docx(p: Path) -> List[str]:
    d = docx.Document(str(p))
    parts = []
    # headings + paragraphs
    for para in d.paragraphs:
        parts.append(para.text)
    # tables
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(norm(cell.text) for cell in row.cells))
    return [norm("\n".join(parts))]


def extract_doc(p: Path) -> List[str]:
    if not mammoth:
        raise RuntimeError("mammoth not installed for .doc. Install requirements.txt")
    with open(p, "rb") as f:
        result = mammoth.extract_raw_text(f)
    return [norm(result.value or "")]


def extract_xlsx(p: Path, max_cells: int = 20000) -> List[str]:
    # Summarize each sheet to readable text with headers
    xls = pd.ExcelFile(p)
    pages = []
    for name in xls.sheet_names:
        df = xls.parse(name)
        # clip very large sheets
        if df.size > max_cells:
            df = df.iloc[: min(1000, len(df)), : 50]
        # coerce to strings, keep headers
        txt = [f"# Sheet: {name}"]
        txt.append(", ".join(map(str, df.columns.tolist())))
        # first N rows only to avoid bloat
        n_show = min(50, len(df))
        for i in range(n_show):
            row = [str(x) for x in df.iloc[i].tolist()]
            txt.append(" | ".join(row))
        pages.append(norm("\n".join(txt)))
    return pages


def extract_images(p: Path) -> List[str]:
    # OCR + basic EXIF summary
    img = Image.open(p)
    ocr = pytesseract.image_to_string(img)
    meta = {}
    try:
        exif = img.getexif()
        if exif:
            meta = {str(k): str(v) for k, v in exif.items() if v}
    except Exception:
        pass
    stamp = [f"# Image: {p.name}", f"# EXIF keys: {', '.join(list(meta.keys())[:10])}", ocr]
    return [norm("\n".join(stamp))]

# ----------------- chunking -----------------

def make_chunks(pieces: List[str], max_tokens: int = 800) -> List[str]:
    budget = max_tokens * 4  # ~4 chars/token heuristic
    chunks, buf, cur = [], [], 0
    for t in pieces:
        t = t or ""
        if cur + len(t) > budget and buf:
            chunks.append(" ".join(buf)); buf.clear(); cur = 0
        buf.append(t); cur += len(t)
    if buf:
        chunks.append(" ".join(buf))
    return [norm(c) for c in chunks if norm(c)]

# ----------------- driver -----------------

HANDLERS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_doc,
    ".xlsx": extract_xlsx,
    ".xls": extract_xlsx,  # openpyxl reads xls via xlrd sometimes; if fails, convert to xlsx
    ".png": extract_images,
    ".jpg": extract_images,
    ".jpeg": extract_images,
    ".tif": extract_images,
    ".tiff": extract_images,
}

from io import BytesIO
import io

def run(cfg_path: str = "config.yaml"):
    import yaml
    cfg = yaml.safe_load(Path(cfg_path).read_text())

    target = cfg["primary"]
    persist = target.get("persist_directory")
    client = chromadb.Client(Settings(persist_directory=persist))
    coll = client.get_or_create_collection(target["collection"], metadata={"gh":"plan"})

    manifest_path = Path(cfg.get("manifest", ".ingest/manifest.json"))
    manifest = json.loads(manifest_path.read_text()) if manifest_path.exists() else {"files": {}}

    # gather files
    include_exts = set([e.lower() for e in cfg.get("extensions", list(HANDLERS.keys()))])
    sources = [Path(p) for p in cfg["sources"]]
    files = []
    for root in sources:
        for p in root.rglob("*"):
            if p.is_file() and p.suffix.lower() in include_exts:
                files.append(p)

    added = skipped = failed = 0
    for p in files:
        try:
            fid = str(p.resolve()); sig = file_sig(p)
            prev = manifest["files"].get(fid)
            if prev and prev.get("hash") == sig:
                skipped += 1; continue

            handler = HANDLERS.get(p.suffix.lower())
            if not handler:
                skipped += 1; continue

            pages = handler(p)
            chunks = make_chunks(pages, cfg.get("chunk_tokens", 800))
            if not chunks:
                skipped += 1; continue

            ids = [sha(f"{fid}::{sig}::{i}".encode()) for i, _ in enumerate(chunks)]
            metas = [{
                "source": fid,
                "file_hash": sig,
                "chunk": i,
                "ext": p.suffix.lower(),
                "ingested_at": int(time.time()),
            } for i, _ in enumerate(chunks)]

            coll.upsert(ids=ids, documents=chunks, metadatas=metas)
            manifest["files"][fid] = {"hash": sig, "chunks": len(ids), "ext": p.suffix.lower()}
            added += 1
        except Exception:
            failed += 1
            traceback.print_exc()

    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, indent=2))

    print(json.dumps({
        "collection": coll.name,
        "added_files": added,
        "skipped_files": skipped,
        "failed_files": failed,
        "total_tracked": len(manifest["files"]),
        "persist_directory": persist
    }, indent=2))

if __name__ == "__main__":
    cfg = sys.argv[1] if len(sys.argv) > 1 else "config.yaml"
    run(cfg)
